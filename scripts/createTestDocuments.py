from pathlib import Path
from zipfile import ZipFile, ZipInfo, ZIP_DEFLATED
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'test'
OUT.mkdir(exist_ok=True)

HEADERS = ['성명', '주민등록번호', '휴대전화', '이메일', '주소', '계좌번호']
ROWS = [
    ['홍길동', '900101-1234568', '010-1234-5678', 'hong.gildong@example.com', '서울특별시 강남구 테헤란로 152 101동 1203호', '110-123-456789'],
    ['김민지', '950505-2234567', '010-9876-5432', 'minji.kim@test.co.kr', '부산광역시 해운대구 센텀중앙로 55 802호', '3333-02-1234567'],
    ['이철수', '850315-1234567', '02-345-6789', 'cs.lee@example.org', '대전광역시 유성구 대학로 99', '1002-345-678901'],
]

EXTRA = [
    '카드번호: 4111-1111-1111-1111',
    '여권번호: M12345678',
    '사업자등록번호: 220-81-62517',
    '차량번호: 12가3456',
    'IP 주소: 192.168.10.25',
    '생년월일: 1990-01-01',
]

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_twips))
    tc_w.set(qn('w:type'), 'dxa')

def create_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.59), Cm(27.94)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)
    sec.header_distance = sec.footer_distance = Cm(1.25)

    normal = doc.styles['Normal']
    normal.font.name = 'Malgun Gothic'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size in [('Heading 1', 16), ('Heading 2', 13)]:
        style = doc.styles[style_name]
        style.font.name = 'Malgun Gothic'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        style.font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run('개인정보 탐지 테스트 문서')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    subtitle = doc.add_paragraph('본 문서의 개인정보는 프로그램 시험용 가상 데이터입니다.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. 민원인 기본정보', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.autofit = False
    table.style = 'Table Grid'
    widths = [900, 1500, 1500, 2100, 2460, 1380]
    for idx, text in enumerate(HEADERS):
        cell = table.rows[0].cells[idx]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, 'DCE6F1')
        set_cell_width(cell, widths[idx])
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.5)
    for row in ROWS:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cells[idx], widths[idx])
            for run in cells[idx].paragraphs[0].runs:
                run.font.size = Pt(8.5)

    doc.add_heading('2. 추가 개인정보 유형', level=1)
    for item in EXTRA:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. 분리된 텍스트 런 시험', level=1)
    p = doc.add_paragraph()
    p.add_run('민원인 홍길')
    p.add_run('동의 연락처는 010-1234-')
    p.add_run('5678이며 이메일은 hong.gildong@example.com입니다.')

    doc.add_heading('4. 머리글 개인정보 시험', level=1)
    doc.add_paragraph('아래 정보는 문서 머리글에도 반복되어 머리글 분석을 확인할 수 있습니다.')
    header = sec.header.paragraphs[0]
    header.text = '담당자 박영희 | 연락처 010-5555-7777 | manager@example.com'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = sec.footer.paragraphs[0]
    footer.text = '개인정보 비식별화 시험용 가상 문서'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT / '개인정보_탐지_테스트.docx')

# Hangul-authored template used as a structural base for the .hwpx fixture.
# A hand-written skeleton (bare header/section, empty manifest, no content.hpf,
# compressed mimetype) is NOT a valid OWPML package and Hangul rejects it with
# "파일이 손상되었습니다". Instead we clone a real Hangul template and swap in only
# section0.xml, referencing char/para/style IDs that already exist in its header.
HWPX_TEMPLATE_GLOBS = [
    r'C:\Program Files (x86)\Hnc\Office*\HOffice*\Shared\HwpTemplate\FrequentComponent\공문서_양식_템플릿_일반.hwpx',
    r'C:\Program Files\Hnc\Office*\HOffice*\Shared\HwpTemplate\FrequentComponent\공문서_양식_템플릿_일반.hwpx',
    r'C:\Program Files*\HNC\Office*\HOffice*\Shared\HwpTemplate\**\*.hwpx',
]


def _find_hwpx_template():
    import glob
    for pattern in HWPX_TEMPLATE_GLOBS:
        for hit in glob.glob(pattern, recursive=True):
            return Path(hit)
    return None


def create_hwpx():
    template = _find_hwpx_template()
    if template is None:
        raise FileNotFoundError(
            'HWPX 템플릿을 찾을 수 없습니다. 한글(HWP)이 설치된 환경에서 실행하거나 '
            'HWPX_TEMPLATE_GLOBS 경로를 조정하세요.')

    paragraphs = [
        '개인정보 탐지 테스트 문서',
        '본 문서의 개인정보는 프로그램 시험용 가상 데이터입니다.',
        '성명 주민등록번호 휴대전화 이메일 주소 계좌번호',
    ]
    paragraphs.extend(' | '.join(row) for row in ROWS)
    paragraphs.extend(EXTRA)
    paragraphs.append('민원인 홍길동의 연락처는 010-1234-5678이며 이메일은 hong.gildong@example.com입니다.')

    with ZipFile(template, 'r') as zin:
        ref_section = zin.read('Contents/section0.xml').decode('utf-8')
        # opening <hs:sec ...> tag carries every namespace declaration we need
        sec_open = re.match(r'^(.*?<hs:sec\b[^>]*>)', ref_section, re.S).group(1)
        # the mandatory page-setup block (page size, margins, note/border props)
        secpr = re.search(r'<hp:secPr\b.*?</hp:secPr>', ref_section, re.S).group(0)

        # IDs known to exist in this template's header.xml
        char_ref, para_ref, style_ref = '29', '1', '0'

        def build_para(text, pid, first=False):
            secpr_xml = secpr if first else ''
            return (
                f'<hp:p id="{pid}" paraPrIDRef="{para_ref}" styleIDRef="{style_ref}" '
                f'pageBreak="0" columnBreak="0" merged="0">'
                f'<hp:run charPrIDRef="{char_ref}">{secpr_xml}'
                f'<hp:t>{escape(text)}</hp:t></hp:run></hp:p>')

        body = ''.join(
            build_para(text, 2000000000 + i, first=(i == 0))
            for i, text in enumerate(paragraphs))
        section = sec_open + body + '</hs:sec>'

        # Rewrite the package verbatim, replacing only section0.xml.
        # Preserving each entry's compression type keeps `mimetype` STORED and
        # first, as the spec requires.
        with ZipFile(OUT / '개인정보_탐지_테스트.hwpx', 'w') as zout:
            for info in zin.infolist():
                data = (section.encode('utf-8')
                        if info.filename == 'Contents/section0.xml'
                        else zin.read(info.filename))
                zi = ZipInfo(info.filename, date_time=info.date_time)
                zi.compress_type = info.compress_type
                zi.external_attr = info.external_attr
                zi.internal_attr = info.internal_attr
                zi.create_system = info.create_system
                zout.writestr(zi, data)

def create_pdf():
    font_path = Path('C:/Windows/Fonts/malgun.ttf')
    if not font_path.exists():
        font_path = Path('C:/Windows/Fonts/gulim.ttc')
    pdfmetrics.registerFont(TTFont('KoreanTestFont', str(font_path)))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('KTitle', parent=styles['Title'], fontName='KoreanTestFont', fontSize=20, leading=25, alignment=TA_CENTER, textColor=colors.HexColor('#1F4D78'))
    body_style = ParagraphStyle('KBody', parent=styles['BodyText'], fontName='KoreanTestFont', fontSize=9, leading=13)
    doc = SimpleDocTemplate(str(OUT / '개인정보_탐지_테스트.pdf'), pagesize=A4, rightMargin=15*mm, leftMargin=15*mm, topMargin=16*mm, bottomMargin=16*mm)
    story = [Paragraph('개인정보 탐지 테스트 문서', title_style), Spacer(1, 5*mm), Paragraph('본 문서의 개인정보는 프로그램 시험용 가상 데이터입니다.', body_style), Spacer(1, 4*mm)]
    data = [[Paragraph(h, body_style) for h in HEADERS]] + [[Paragraph(v, body_style) for v in row] for row in ROWS]
    table = Table(data, colWidths=[18*mm, 28*mm, 27*mm, 39*mm, 54*mm, 28*mm], repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#DCE6F1')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#17365D')),
        ('FONTNAME', (0,0), (-1,-1), 'KoreanTestFont'),
        ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor('#A6B7C8')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 3), ('RIGHTPADDING', (0,0), (-1,-1), 3),
        ('TOPPADDING', (0,0), (-1,-1), 4), ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))
    story.extend([table, Spacer(1, 5*mm), Paragraph('추가 개인정보 유형', title_style)])
    for item in EXTRA:
        story.append(Paragraph('• ' + item, body_style))
    story.extend([Spacer(1, 3*mm), Paragraph('반복 문장: 민원인 홍길동에게 010-1234-5678로 연락하고 hong.gildong@example.com으로 결과를 전송합니다.', body_style)])
    doc.build(story)

def normalize_xlsx_for_safedoc():
    """artifact-tool의 x: 접두사/str 셀을 일반 OOXML inlineStr 표현으로 정규화한다."""
    source = OUT / '개인정보_탐지_테스트.xlsx'
    if not source.exists():
        return
    temp = source.with_suffix('.xlsx.tmp')
    with ZipFile(source, 'r') as zin, ZipFile(temp, 'w', ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == 'xl/workbook.xml' or info.filename == 'xl/sharedStrings.xml' or info.filename.startswith('xl/worksheets/sheet'):
                xml = data.decode('utf-8')
                xml = xml.replace('xmlns:x=', 'xmlns=').replace('<x:', '<').replace('</x:', '</')
                if info.filename.startswith('xl/worksheets/sheet'):
                    xml = re.sub(
                        r'<c([^>]*)\s+t="str"><v>([\s\S]*?)</v></c>',
                        r'<c\1 t="inlineStr"><is><t>\2</t></is></c>',
                        xml,
                    )
                data = xml.encode('utf-8')
            zout.writestr(info, data)
    temp.replace(source)

if __name__ == '__main__':
    normalize_xlsx_for_safedoc()
    create_docx()
    create_hwpx()
    create_pdf()
    for name in ['개인정보_탐지_테스트.docx', '개인정보_탐지_테스트.hwpx', '개인정보_탐지_테스트.pdf']:
        path = OUT / name
        print(f'{path.name}: {path.stat().st_size} bytes')
